"""
Report generation module for StockMate.

Provides advanced filtering, custom expressions, and multi-format export
(Excel, PDF, Word) for inventory data using pandas, reportlab, and
python-docx.
"""

import re
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
)

# ---------------------------------------------------------------------------
# Operator implementations — pure, predictable, one responsibility each
# ---------------------------------------------------------------------------

_OPERATOR_MAP: dict[str, str] = {
    "equals": "==",
    "contains": ".str.contains({value}, na=False, regex=False)",
    "gt": ">",
    "lt": "<",
    "gte": ">=",
    "lte": "<=",
    "starts_with": ".str.startswith({value}, na=False)",
    "ends_with": ".str.endswith({value}, na=False)",
    "regex": ".str.contains({value}, na=True, regex=True)",
}


def _build_mask(
    df: pd.DataFrame,
    field: str,
    operator: str,
    value: Any,
) -> pd.Series:
    """Build a boolean mask for a single filter condition.

    Returns a Series of bools aligned to *df*.
    """
    if operator == "is_empty":
        return df[field].isna() | (df[field].astype(str).str.strip() == "")

    if operator == "not_empty":
        return df[field].notna() & (df[field].astype(str).str.strip() != "")

    if operator == "modulo":
        numeric = pd.to_numeric(df[field], errors="coerce")
        return (numeric % value) == 0

    if operator in ("contains", "regex"):
        pattern = str(value)
        if operator == "contains":
            return df[field].astype(str).str.contains(pattern, na=False, regex=False)
        return df[field].astype(str).str.contains(pattern, na=True, regex=True)

    if operator in ("starts_with", "ends_with"):
        target = str(value)
        if operator == "starts_with":
            return df[field].astype(str).str.startswith(target, na=False)
        return df[field].astype(str).str.endswith(target, na=False)

    if operator == "equals":
        return df[field] == value

    # Numeric comparisons: gt, lt, gte, lte
    numeric_col = pd.to_numeric(df[field], errors="coerce")
    numeric_val = float(value)

    if operator == "gt":
        return numeric_col > numeric_val
    if operator == "lt":
        return numeric_col < numeric_val
    if operator == "gte":
        return numeric_col >= numeric_val
    if operator == "lte":
        return numeric_col <= numeric_val

    raise ValueError(f"Unknown filter operator: {operator!r}")


def _combine_masks(
    base: pd.Series,
    new_mask: pd.Series,
    logic: str,
) -> pd.Series:
    """Combine two boolean masks using the specified logic gate."""
    if logic == "AND":
        return base & new_mask
    if logic == "OR":
        return base | new_mask
    if logic == "AND NOT":
        return base & ~new_mask
    if logic == "OR NOT":
        return base | ~new_mask
    if logic == "XOR":
        return base ^ new_mask
    raise ValueError(f"Unknown logic gate: {logic!r}")


# ---------------------------------------------------------------------------
# ReportGenerator
# ---------------------------------------------------------------------------


class ReportGenerator:
    """Filter, transform, and export inventory data in multiple formats."""

    def __init__(self, inventory_manager: Any) -> None:
        self._inventory = inventory_manager

    # ------------------------------------------------------------------
    # Filtering
    # ------------------------------------------------------------------

    def apply_filters(
        self,
        df: pd.DataFrame,
        conditions: list[dict[str, Any]],
    ) -> pd.DataFrame:
        """Apply a list of filter conditions with logic gates.

        Each condition dict must contain:
        - ``logic``: one of ``"AND"``, ``"OR"``, ``"AND NOT"``, ``"OR NOT"``,
          ``"XOR"``.  The first condition ignores this key (acts as base).
        - ``field``: column name in *df*.
        - ``operator``: one of the supported operators.
        - ``value``: the value to compare against (unused for ``is_empty``,
          ``not_empty``).

        Returns a filtered DataFrame.
        """
        if not conditions:
            return df

        combined_mask: pd.Series | None = None

        for i, cond in enumerate(conditions):
            field = cond["field"]
            operator = cond["operator"]
            value = cond.get("value")
            logic = cond.get("logic", "AND")

            if field not in df.columns:
                continue

            mask = _build_mask(df, field, operator, value)

            if i == 0 or combined_mask is None:
                combined_mask = mask
            else:
                combined_mask = _combine_masks(combined_mask, mask, logic)

        if combined_mask is None:
            return df

        return df[combined_mask].copy()

    def apply_custom_expression(
        self,
        df: pd.DataFrame,
        expression: str,
    ) -> pd.DataFrame:
        """Filter *df* using a pandas ``query()`` expression string.

        Returns the filtered DataFrame.  Raises ``ValueError`` on invalid
        expressions so the caller can surface a clear error to the user.
        """
        if not expression or not expression.strip():
            return df

        try:
            return df.query(expression).copy()
        except Exception as exc:
            raise ValueError(
                f"Invalid query expression: {expression!r} — {exc}"
            ) from exc

    def apply_limit(
        self,
        df: pd.DataFrame,
        limit: int | None = None,
        modulo: int | None = None,
    ) -> pd.DataFrame:
        """Apply row limiting and/or modulo-based row selection.

        - *limit*: keep only the first N rows.
        - *modulo*: keep rows whose 0-based index satisfies ``index % modulo == 0``.
        """
        result = df

        if modulo is not None and modulo > 0:
            result = result.iloc[::modulo]

        if limit is not None and limit > 0:
            result = result.head(limit)

        return result.copy()

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------

    def export(
        self,
        df: pd.DataFrame,
        format: str,
        filepath: str,
    ) -> bool:
        """Export *df* to the specified *format* at *filepath*.

        Supported formats: ``"excel"``, ``"pdf"``, ``"word"``.
        Returns ``True`` on success, ``False`` on failure.
        """
        fmt = format.lower().strip()

        if fmt == "excel":
            return self._export_excel(df, filepath)
        if fmt == "pdf":
            return self._export_pdf(df, filepath)
        if fmt == "word":
            return self._export_word(df, filepath)

        raise ValueError(f"Unsupported export format: {format!r}")

    @staticmethod
    def _export_excel(df: pd.DataFrame, filepath: str) -> bool:
        """Write DataFrame to an Excel file."""
        try:
            df.to_excel(filepath, index=False)
            return True
        except Exception:
            return False

    @staticmethod
    def _export_pdf(df: pd.DataFrame, filepath: str) -> bool:
        """Write DataFrame to a PDF document using ReportLab."""
        try:
            doc = SimpleDocTemplate(
                filepath,
                pagesize=letter,
                leftMargin=0.5 * inch,
                rightMargin=0.5 * inch,
            )
            elements: list = []

            styles = getSampleStyleSheet()
            elements.append(Paragraph("Inventory Report", styles["Title"]))
            elements.append(Spacer(1, 12))

            # Prepare table data
            headers = [str(col) for col in df.columns]
            rows = [headers]
            for _, row in df.iterrows():
                rows.append([str(v) for v in row.values])

            col_widths = [
                min(2.0 * inch, max(0.6 * inch, letter[0] / len(headers)))
                for _ in headers
            ]

            table = Table(rows, colWidths=col_widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#007acc")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, 0), 8),
                        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                        ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                        ("TEXTCOLOR", (0, 1), (-1, -1), colors.black),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 1), (-1, -1), 7),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        (
                            "ROWBACKGROUNDS",
                            (0, 1),
                            (-1, -1),
                            [colors.white, colors.HexColor("#f0f4f8")],
                        ),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]
                )
            )

            elements.append(table)
            doc.build(elements)
            return True
        except Exception:
            return False

    @staticmethod
    def _export_word(df: pd.DataFrame, filepath: str) -> bool:
        """Write DataFrame to a Word document using python-docx."""
        try:
            doc = Document()
            doc.add_heading("Inventory Report", level=1)

            num_rows = len(df) + 1  # +1 for header
            num_cols = len(df.columns)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = "Light Grid Accent 1"

            # Header row
            for col_idx, col_name in enumerate(df.columns):
                cell = table.rows[0].cells[col_idx]
                cell.text = str(col_name)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)

            # Data rows
            for row_idx, (_, row_data) in enumerate(df.iterrows(), start=1):
                for col_idx, value in enumerate(row_data.values):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

            doc.save(filepath)
            return True
        except Exception:
            return False
